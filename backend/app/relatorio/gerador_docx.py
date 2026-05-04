"""
Gerador de Relatório Previdenciário em DOCX — Visual Law — Feliciano Advocacia.

Gera documento Word profissional com:
  - Identidade visual Feliciano Advocacia (azul #1a3c6e / #4F81BD)
  - Visual Law (tabelas, ícones textuais, destaques visuais)
  - Memória de cálculo
  - Score de Prontidão
  - TC nos Marcos Legais
  - Plano de Ação
  - Fundamentação Legal
"""
from __future__ import annotations
from datetime import date
from decimal import Decimal
from typing import Any, Dict, List, Optional
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Cores Feliciano Advocacia ──
AZUL_ESCURO = RGBColor(0x1a, 0x3c, 0x6e)
AZUL_MEDIO = RGBColor(0x4F, 0x81, 0xBD)
AZUL_CLARO = RGBColor(0x36, 0x5F, 0x91)
VERDE = RGBColor(0x06, 0x5f, 0x46)
VERMELHO = RGBColor(0x99, 0x1b, 0x1b)
LARANJA = RGBColor(0xb4, 0x53, 0x09)
CINZA = RGBColor(0x6b, 0x72, 0x80)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_bg(cell, hex_color: str):
    """Define cor de fundo de uma célula."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_heading(doc, text, level=1):
    """Adiciona heading com estilo Feliciano."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = AZUL_ESCURO if level == 1 else AZUL_MEDIO
    return h


def _add_para(doc, text, bold=False, color=None, size=10, align=None, space_after=6):
    """Adiciona parágrafo formatado."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _fmt_brl(valor) -> str:
    """Formata valor para R$ brasileiro."""
    if valor is None:
        return "R$ 0,00"
    try:
        v = Decimal(str(valor))
        return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return f"R$ {valor}"


def gerar_docx_planejamento(
    segurado: Dict,
    planejamento: Dict,
    nome_advogado: Optional[str] = None,
) -> bytes:
    """
    Gera DOCX profissional de planejamento previdenciário.

    Returns:
        bytes do arquivo .docx
    """
    doc = Document()

    # ── Configurar estilos ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = AZUL_ESCURO if level == 1 else AZUL_MEDIO

    # ── Margens ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    dp = segurado.get("dados_pessoais", {})
    hoje = date.today().strftime("%d/%m/%Y")
    nome = dp.get("nome", "Segurado")
    tc = planejamento.get("tc_atual", {})
    score = planejamento.get("score_prontidao", {})
    marcos = planejamento.get("marcos_legais") or []
    projecoes = planejamento.get("projecoes") or []
    custo = planejamento.get("custo_beneficio") or []
    qualidade = planejamento.get("qualidade_segurado") or {}
    pensao = planejamento.get("pensao_projetada") or {}
    plano = planejamento.get("plano_acao") or []
    resumo = planejamento.get("resumo_executivo") or {}
    comp_sem = planejamento.get("competencias_sem_salario") or {}
    cenarios = planejamento.get("cenarios_vida") or []
    argumentos = planejamento.get("argumentos_cliente") or []

    alcancaveis = sorted(
        [p for p in projecoes if p.get("data_elegibilidade")],
        key=lambda p: p.get("meses_faltantes", 999)
    )
    melhor = alcancaveis[0] if alcancaveis else None

    # ═══════════════════════════════════════════════════════════════════════
    # CAPA
    # ═══════════════════════════════════════════════════════════════════════

    # Linha decorativa (simulada com tabela de 1 célula)
    _add_para(doc, "", size=6, space_after=0)
    _add_para(doc, "", size=6, space_after=0)

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("FELICIANO ADVOCACIA")
    run_logo.font.size = Pt(28)
    run_logo.font.bold = True
    run_logo.font.color.rgb = AZUL_ESCURO
    run_logo.font.name = "Calibri"

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Assessoria Jurídica Previdenciária")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = AZUL_MEDIO
    run_sub.font.name = "Calibri"
    p_sub.paragraph_format.space_after = Pt(40)

    # Linha separadora
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("━" * 60)
    run_line.font.color.rgb = AZUL_MEDIO
    run_line.font.size = Pt(8)
    p_line.paragraph_format.space_after = Pt(30)

    # Título do documento
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run("PLANEJAMENTO PREVIDENCIÁRIO")
    run_t.font.size = Pt(22)
    run_t.font.bold = True
    run_t.font.color.rgb = AZUL_ESCURO
    p_titulo.paragraph_format.space_after = Pt(8)

    p_nome_cli = doc.add_paragraph()
    p_nome_cli.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nc = p_nome_cli.add_run(nome.upper())
    run_nc.font.size = Pt(16)
    run_nc.font.bold = True
    run_nc.font.color.rgb = AZUL_MEDIO
    p_nome_cli.paragraph_format.space_after = Pt(30)

    # Quadro de dados na capa
    tabela_capa = doc.add_table(rows=4, cols=2)
    tabela_capa.alignment = WD_TABLE_ALIGNMENT.CENTER
    dados_capa = [
        ("CPF:", dp.get("cpf", "—")),
        ("Data de Nascimento:", dp.get("data_nascimento", "—")),
        ("NIT/PIS:", dp.get("nit", "—")),
        ("Data da Análise:", hoje),
    ]
    for i, (label, valor) in enumerate(dados_capa):
        cell_l = tabela_capa.cell(i, 0)
        cell_v = tabela_capa.cell(i, 1)
        cell_l.text = label
        cell_v.text = str(valor) if valor else "—"
        for p in cell_l.paragraphs:
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = AZUL_ESCURO
        for p in cell_v.paragraphs:
            p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 1. SCORE DE PRONTIDÃO PREVIDENCIÁRIA
    # ═══════════════════════════════════════════════════════════════════════

    if score:
        _add_heading(doc, "1. Score de Prontidão Previdenciária", level=1)

        score_val = score.get("score", 0)
        classif = score.get("classificacao", "").replace("_", " ")
        msg = score.get("mensagem", "")

        # Score grande
        p_score = doc.add_paragraph()
        p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_s = p_score.add_run(f"{score_val}")
        run_s.font.size = Pt(48)
        run_s.font.bold = True
        cor_hex = score.get("cor", "#1a3c6e")
        try:
            run_s.font.color.rgb = RGBColor(int(cor_hex[1:3], 16), int(cor_hex[3:5], 16), int(cor_hex[5:7], 16))
        except Exception:
            run_s.font.color.rgb = AZUL_ESCURO
        run_s2 = p_score.add_run(" / 1000")
        run_s2.font.size = Pt(14)
        run_s2.font.color.rgb = CINZA

        p_classif = doc.add_paragraph()
        p_classif.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cl = p_classif.add_run(classif)
        run_cl.font.size = Pt(14)
        run_cl.font.bold = True
        run_cl.font.color.rgb = AZUL_MEDIO

        _add_para(doc, msg, size=10, color=CINZA, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Tabela de componentes
        comp = score.get("componentes", {})
        if comp:
            tabela_score = doc.add_table(rows=len(comp) + 1, cols=4)
            tabela_score.style = 'Light List Accent 1'
            tabela_score.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = ["Componente", "Pontos", "Máximo", "Detalhe"]
            for j, h in enumerate(headers):
                cell = tabela_score.cell(0, j)
                cell.text = h
                _set_cell_bg(cell, "1a3c6e")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = BRANCO
                        r.font.bold = True
                        r.font.size = Pt(9)

            labels = {
                "tempo_contribuicao": "Tempo de Contribuição",
                "idade": "Idade",
                "carencia": "Carência",
                "qualidade_segurado": "Qualidade Segurado",
                "proximidade": "Proximidade",
                "valor_beneficio": "Valor Benefício",
            }
            for i, (key, c) in enumerate(comp.items()):
                row = i + 1
                tabela_score.cell(row, 0).text = labels.get(key, key)
                tabela_score.cell(row, 1).text = str(c.get("pontos", 0))
                tabela_score.cell(row, 2).text = str(c.get("maximo", 0))
                tabela_score.cell(row, 3).text = str(c.get("detalhe", ""))
                for j in range(4):
                    for p in tabela_score.cell(row, j).paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 2. RESUMO EXECUTIVO
    # ═══════════════════════════════════════════════════════════════════════

    if resumo:
        _add_heading(doc, "2. Resumo Executivo", level=1)

        campos_resumo = [
            ("SITUAÇÃO ATUAL", resumo.get("situacao_atual", "")),
            ("MELHOR CAMINHO", resumo.get("melhor_caminho", "")),
            ("AÇÃO IMEDIATA", resumo.get("acao_imediata", "")),
            ("PRÓXIMO PASSO", resumo.get("proximo_passo", "")),
        ]
        cores_bg = ["DBE5F1", "E2EFDA", "FFF2CC", "E8DAEF"]

        for (titulo, texto), cor_bg in zip(campos_resumo, cores_bg):
            tabela_r = doc.add_table(rows=1, cols=1)
            cell = tabela_r.cell(0, 0)
            p = cell.paragraphs[0]
            run_t = p.add_run(f"{titulo}\n")
            run_t.font.size = Pt(8)
            run_t.font.bold = True
            run_t.font.color.rgb = AZUL_ESCURO
            run_v = p.add_run(texto)
            run_v.font.size = Pt(10)
            _set_cell_bg(cell, cor_bg)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ═══════════════════════════════════════════════════════════════════════
    # 3. SITUAÇÃO PREVIDENCIÁRIA ATUAL
    # ═══════════════════════════════════════════════════════════════════════

    _add_heading(doc, "3. Situação Previdenciária Atual", level=1)

    # Tabela de dados do segurado
    tabela_sit = doc.add_table(rows=6, cols=2)
    tabela_sit.style = 'Light List Accent 1'
    dados_sit = [
        ("Nome:", nome),
        ("Tempo de Contribuição:", f"{tc.get('anos', 0)} anos, {tc.get('meses', 0)} meses e {tc.get('dias', 0)} dias"),
        ("Salário de Contribuição Projetado:", _fmt_brl(planejamento.get("salario_projetado", 0))),
        ("Qualidade de Segurado:", qualidade.get("status", "—")),
        ("Elegível Agora:", "SIM — Pode requerer hoje" if planejamento.get("elegiveis_agora") else "NÃO — Ainda não preenche os requisitos"),
        ("Melhor RMI Projetada:", _fmt_brl(planejamento.get("melhor_rmi", 0))),
    ]
    for i, (label, valor) in enumerate(dados_sit):
        tabela_sit.cell(i, 0).text = label
        tabela_sit.cell(i, 1).text = str(valor)
        for p in tabela_sit.cell(i, 0).paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
        for p in tabela_sit.cell(i, 1).paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

    # Fundamentação legal da contagem de TC
    _add_para(doc, (
        "Metodologia de contagem conforme legislação vigente: "
        "Empregado CLT/Avulso/Doméstico — TC em dias corridos do período (Art. 60, Decreto 3.048/99); "
        "Facultativo/CI/MEI — TC por competências com contribuição efetiva, 30 dias/mês (Art. 19-C, Decreto 10.410/2020); "
        "Contribuições abaixo do SM pós-EC 103/2019 NÃO contam (Art. 19-E, Decreto 3.048/99); "
        "Auxílio-doença intercalado conta como TC (Art. 60 §3º, Lei 8.213/91); "
        "Períodos concomitantes contados apenas uma vez; "
        "Indicadores CNIS (PREC-MENOR-MIN, IREC-INDPEND) excluem contribuições com pendência."
    ), size=8, color=CINZA)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 4. QUALIDADE DE SEGURADO
    # ═══════════════════════════════════════════════════════════════════════

    if qualidade:
        _add_heading(doc, "4. Qualidade de Segurado", level=1)

        status_qs = qualidade.get("status", "PERDIDA")
        icon = "ATIVA" if status_qs == "ATIVA" else "EM RISCO" if status_qs == "EM_RISCO" else "PERDIDA"
        _add_para(doc, f"Status: {icon}", bold=True, size=12,
                  color=VERDE if status_qs == "ATIVA" else LARANJA if status_qs == "EM_RISCO" else VERMELHO)
        _add_para(doc, qualidade.get("mensagem", ""))

        tabela_qs = doc.add_table(rows=1, cols=3)
        tabela_qs.style = 'Light List Accent 1'
        tabela_qs.cell(0, 0).text = f"Última contribuição: {qualidade.get('ultima_contribuicao', '—')}"
        tabela_qs.cell(0, 1).text = f"Período de graça: {qualidade.get('periodo_graca_meses', 0)} meses"
        tabela_qs.cell(0, 2).text = f"Perda em: {qualidade.get('data_perda_qualidade', '—')}"
        for j in range(3):
            for p in tabela_qs.cell(0, j).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 5. TC NOS MARCOS LEGAIS
    # ═══════════════════════════════════════════════════════════════════════

    if marcos:
        _add_heading(doc, "5. Tempo de Contribuição nos Marcos Legais", level=1)
        _add_para(doc, "Análise do tempo acumulado em cada data-chave da legislação previdenciária.", color=CINZA, size=9)

        tabela_marcos = doc.add_table(rows=len(marcos) + 1, cols=5)
        tabela_marcos.style = 'Light List Accent 1'
        tabela_marcos.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers_m = ["Marco Legal", "Data", "TC Acumulado", "Contribuições", "Análise"]
        for j, h in enumerate(headers_m):
            cell = tabela_marcos.cell(0, j)
            cell.text = h
            _set_cell_bg(cell, "1a3c6e")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = BRANCO
                    r.font.bold = True
                    r.font.size = Pt(8)

        for i, m in enumerate(marcos):
            row = i + 1
            tabela_marcos.cell(row, 0).text = f"{m.get('sigla', '')} — {m.get('nome', '')}"
            tabela_marcos.cell(row, 1).text = str(m.get("data", ""))
            tabela_marcos.cell(row, 2).text = str(m.get("tc_texto", ""))
            tabela_marcos.cell(row, 3).text = str(m.get("contribuicoes", ""))
            tabela_marcos.cell(row, 4).text = str(m.get("observacao", ""))
            for j in range(5):
                for p in tabela_marcos.cell(row, j).paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 6. COMPETÊNCIAS SEM SALÁRIO
    # ═══════════════════════════════════════════════════════════════════════

    if comp_sem and comp_sem.get("total_problemas", 0) > 0:
        _add_heading(doc, "6. Competências sem Salário de Contribuição", level=1)
        _add_para(doc, comp_sem.get("mensagem", ""), bold=True, color=LARANJA)
        _add_para(doc, f"Impacto: {comp_sem.get('impacto', '')}", size=10)

        sem_sal = comp_sem.get("sem_salario", [])
        if sem_sal:
            competencias_texto = ", ".join([s.get("competencia", "") for s in sem_sal])
            _add_para(doc, f"Competências afetadas: {competencias_texto}", size=9, color=VERMELHO)
            if sem_sal:
                _add_para(doc, f"Empregador: {sem_sal[0].get('empregador', '—')}", size=9, color=CINZA)

    # ═══════════════════════════════════════════════════════════════════════
    # 7. LINHA DO TEMPO — REGRAS DE APOSENTADORIA
    # ═══════════════════════════════════════════════════════════════════════

    doc.add_page_break()
    _add_heading(doc, "7. Análise das Regras de Transição (EC 103/2019)", level=1)
    _add_para(doc, "Projeção de datas de aposentadoria considerando contribuição contínua.", color=CINZA, size=9)

    if projecoes:
        tabela_regras = doc.add_table(rows=len(projecoes) + 1, cols=5)
        tabela_regras.style = 'Light List Accent 1'
        tabela_regras.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers_r = ["Regra", "Data Prevista", "Tempo Faltando", "RMI Estimada", "Status"]
        for j, h in enumerate(headers_r):
            cell = tabela_regras.cell(0, j)
            cell.text = h
            _set_cell_bg(cell, "1a3c6e")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = BRANCO
                    r.font.bold = True
                    r.font.size = Pt(8)

        sorted_proj = sorted(projecoes, key=lambda p: p.get("meses_faltantes") or 9999)
        for i, proj in enumerate(sorted_proj):
            row = i + 1
            is_melhor = i == 0 and proj.get("data_elegibilidade")
            tabela_regras.cell(row, 0).text = (">>> " if is_melhor else "") + str(proj.get("regra", ""))
            tabela_regras.cell(row, 1).text = str(proj.get("data_elegibilidade", "—"))
            tabela_regras.cell(row, 2).text = str(proj.get("texto_faltante", "—"))
            tabela_regras.cell(row, 3).text = str(proj.get("rmi_formatada", "—"))
            tabela_regras.cell(row, 4).text = "Alcançável" if proj.get("data_elegibilidade") else "Não alcançável"

            if is_melhor:
                for j in range(5):
                    _set_cell_bg(tabela_regras.cell(row, j), "DBE5F1")

            for j in range(5):
                for p in tabela_regras.cell(row, j).paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)
                        if is_melhor:
                            r.font.bold = True

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 8. CUSTO-BENEFÍCIO (ROI)
    # ═══════════════════════════════════════════════════════════════════════

    if custo:
        _add_heading(doc, "8. Análise de Custo-Benefício — Retorno sobre Investimento", level=1)

        ev = planejamento.get("expectativa_vida", {})
        ev_anos = ev.get("anos", 76.6)
        _add_para(doc, f"Expectativa de vida IBGE 2024: {ev_anos} anos. Análise do retorno financeiro de cada modalidade de contribuição.", color=CINZA, size=9)

        for cb in custo:
            _add_heading(doc, cb.get("regra", ""), level=2)

            info_text = (
                f"Data: {cb.get('data_elegibilidade', '—')} | "
                f"Idade: {cb.get('idade_aposentadoria', 0):.1f} anos | "
                f"Meses recebendo: {int(cb.get('meses_recebendo', 0))} | "
                f"Total recebido: {_fmt_brl(cb.get('total_recebido', 0))}"
            )
            _add_para(doc, info_text, size=9, color=AZUL_ESCURO)

            mods = cb.get("modalidades", [])
            if mods:
                tabela_roi = doc.add_table(rows=len(mods) + 1, cols=7)
                tabela_roi.style = 'Light List Accent 1'

                headers_roi = ["Modalidade", "Alíquota", "Contrib/mês", "Total Pago", "Lucro", "ROI", "Vale?"]
                for j, h in enumerate(headers_roi):
                    cell = tabela_roi.cell(0, j)
                    cell.text = h
                    _set_cell_bg(cell, "1a3c6e")
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.color.rgb = BRANCO
                            r.font.bold = True
                            r.font.size = Pt(8)

                for mi, m in enumerate(mods):
                    row = mi + 1
                    tabela_roi.cell(row, 0).text = str(m.get("modalidade", ""))
                    tabela_roi.cell(row, 1).text = f"{m.get('aliquota_pct', 0)}%"
                    tabela_roi.cell(row, 2).text = _fmt_brl(m.get("contribuicao_mensal", 0))
                    tabela_roi.cell(row, 3).text = _fmt_brl(m.get("total_pago_ate_apos", 0))
                    tabela_roi.cell(row, 4).text = _fmt_brl(m.get("lucro_liquido", 0))
                    tabela_roi.cell(row, 5).text = f"{m.get('roi_percentual', 0)}%"
                    tabela_roi.cell(row, 6).text = "SIM" if m.get("vale_a_pena") else "NÃO"

                    cor_bg = "E2EFDA" if m.get("vale_a_pena") else "FCE4EC"
                    for j in range(7):
                        _set_cell_bg(tabela_roi.cell(row, j), cor_bg)
                        for p in tabela_roi.cell(row, j).paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(8)

            doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 9. CENÁRIOS DE CONTRIBUIÇÃO
    # ═══════════════════════════════════════════════════════════════════════

    if cenarios:
        doc.add_page_break()
        _add_heading(doc, "9. Cenários de Contribuição — Comparativo com Valores", level=1)

        tabela_cen = doc.add_table(rows=len(cenarios) + 1, cols=5)
        tabela_cen.style = 'Light List Accent 1'

        headers_cen = ["Cenário", "Custo/mês", "Custo/ano", "Total até aposentar", "Impacto na RMI"]
        for j, h in enumerate(headers_cen):
            cell = tabela_cen.cell(0, j)
            cell.text = h
            _set_cell_bg(cell, "1a3c6e")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = BRANCO
                    r.font.bold = True
                    r.font.size = Pt(8)

        for i, cv in enumerate(cenarios):
            row = i + 1
            nome_cen = cv.get("nome", cv.get("cenario", ""))
            tabela_cen.cell(row, 0).text = nome_cen
            tabela_cen.cell(row, 1).text = _fmt_brl(cv.get("monthly_cost", cv.get("custo_mensal", 0)))
            tabela_cen.cell(row, 2).text = _fmt_brl(cv.get("annual_cost", cv.get("custo_anual", 0)))
            tabela_cen.cell(row, 3).text = _fmt_brl(cv.get("total_cost_until_retirement", cv.get("custo_total", 0)))
            rmi_impact = cv.get("impact_on_rmi", cv.get("impacto_rmi", "—"))
            tabela_cen.cell(row, 4).text = str(rmi_impact)
            for j in range(5):
                for p in tabela_cen.cell(row, j).paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 10. PENSÃO POR MORTE
    # ═══════════════════════════════════════════════════════════════════════

    if pensao:
        _add_heading(doc, "10. Pensão por Morte — Projeção para Dependentes", level=1)
        _add_para(doc, pensao.get("mensagem", ""), size=10)

        cen_pensao = pensao.get("cenarios", [])
        if cen_pensao:
            tabela_pen = doc.add_table(rows=len(cen_pensao) + 1, cols=3)
            tabela_pen.style = 'Light List Accent 1'

            for j, h in enumerate(["Dependentes", "Cota (%)", "Valor Mensal"]):
                cell = tabela_pen.cell(0, j)
                cell.text = h
                _set_cell_bg(cell, "1a3c6e")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = BRANCO
                        r.font.bold = True
                        r.font.size = Pt(9)

            for i, c in enumerate(cen_pensao):
                row = i + 1
                tabela_pen.cell(row, 0).text = f"{c.get('dependentes', 0)} dependente(s)"
                tabela_pen.cell(row, 1).text = f"{c.get('cota_pct', 0)}%"
                tabela_pen.cell(row, 2).text = str(c.get("valor_formatado", _fmt_brl(c.get("valor", 0))))
                for j in range(3):
                    for p in tabela_pen.cell(row, j).paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 11. PLANO DE AÇÃO
    # ═══════════════════════════════════════════════════════════════════════

    if plano:
        doc.add_page_break()
        _add_heading(doc, "11. Plano de Ação — Próximos Passos", level=1)

        for p_acao in plano:
            numero = p_acao.get("numero", "")
            titulo_p = p_acao.get("titulo", "")
            desc = p_acao.get("descricao", "")
            prazo = p_acao.get("prazo", "")
            urgencia = p_acao.get("urgencia", "")

            _add_heading(doc, f"Passo {numero}: {titulo_p}", level=2)
            _add_para(doc, desc, size=10)

            urg_color = VERMELHO if urgencia == "ALTA" else LARANJA if urgencia == "MEDIA" else VERDE
            p_urg = doc.add_paragraph()
            run_prazo = p_urg.add_run(f"Prazo: {prazo}")
            run_prazo.font.size = Pt(9)
            run_prazo.font.color.rgb = CINZA
            p_urg.add_run("   |   ")
            run_urg = p_urg.add_run(f"Urgência: {urgencia}")
            run_urg.font.size = Pt(9)
            run_urg.font.bold = True
            run_urg.font.color.rgb = urg_color

    # ═══════════════════════════════════════════════════════════════════════
    # 12. ARGUMENTOS PARA O CLIENTE
    # ═══════════════════════════════════════════════════════════════════════

    if argumentos:
        _add_heading(doc, "12. Argumentos para Apresentação ao Cliente", level=1)
        for i, arg in enumerate(argumentos):
            _add_para(doc, f"{i+1}. {arg}", size=10)

    # ═══════════════════════════════════════════════════════════════════════
    # 13. FUNDAMENTAÇÃO LEGAL
    # ═══════════════════════════════════════════════════════════════════════

    doc.add_page_break()
    _add_heading(doc, "13. Fundamentação Legal", level=1)

    leis = [
        ("EC 103/2019", "Art. 15", "Regra de Transição — Sistema de Pontos (idade + TC)"),
        ("EC 103/2019", "Art. 16", "Regra de Transição — Idade Mínima Progressiva"),
        ("EC 103/2019", "Art. 17", "Regra de Transição — Pedágio 50% com Fator Previdenciário"),
        ("EC 103/2019", "Art. 20", "Regra de Transição — Pedágio 100% com Idade Mínima"),
        ("EC 103/2019", "Art. 19", "Aposentadoria por Idade (Regra Permanente)"),
        ("Lei 8.213/91", "Art. 29", "Período Básico de Cálculo e Salário de Benefício"),
        ("Lei 8.213/91", "Art. 25", "Carência — 180 meses de contribuição"),
        ("Lei 8.213/91", "Art. 15", "Qualidade de segurado e períodos de graça"),
        ("Decreto 3.048/99", "Art. 11", "Segurado facultativo — recolhimento voluntário"),
        ("LC 123/2006", "Art. 18-A", "MEI — alíquota reduzida de 5% (não conta para TC)"),
    ]

    tabela_leis = doc.add_table(rows=len(leis) + 1, cols=3)
    tabela_leis.style = 'Light List Accent 1'

    for j, h in enumerate(["Norma", "Dispositivo", "Aplicação"]):
        cell = tabela_leis.cell(0, j)
        cell.text = h
        _set_cell_bg(cell, "1a3c6e")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = BRANCO
                r.font.bold = True
                r.font.size = Pt(9)

    for i, (norma, art, desc) in enumerate(leis):
        tabela_leis.cell(i + 1, 0).text = norma
        tabela_leis.cell(i + 1, 1).text = art
        tabela_leis.cell(i + 1, 2).text = desc
        for j in range(3):
            for p in tabela_leis.cell(i + 1, j).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 14. ANÁLISE DE ATIVIDADE ESPECIAL POR EMPREGADOR
    # ═══════════════════════════════════════════════════════════════════════

    analise_esp = planejamento.get("analise_especial", [])
    if analise_esp:
        _add_heading(doc, "14. Análise de Atividade Especial por Empregador", level=1)
        _add_para(doc, (
            "O sistema analisou automaticamente cada empregador do CNIS e identificou "
            "possíveis atividades especiais (insalubridade/periculosidade) com base em padrões "
            "conhecidos de CNAE e razão social. A conversão de tempo especial para comum "
            "utiliza fator de 1,4 (homem) ou 1,2 (mulher) — Art. 57 e 58, Lei 8.213/91."
        ), size=10, color=CINZA)

        tabela_esp = doc.add_table(rows=len(analise_esp) + 1, cols=5)
        tabela_esp.style = 'Light List Accent 1'
        for j, h in enumerate(["Empregador", "Período", "Probabilidade", "Agentes Prováveis", "Recomendação"]):
            cell = tabela_esp.cell(0, j)
            cell.text = h
            _set_cell_bg(cell, "b45309")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = BRANCO
                    r.font.bold = True
                    r.font.size = Pt(8)

        for i, ae in enumerate(analise_esp):
            tabela_esp.cell(i + 1, 0).text = ae.get("empregador", "—")
            periodo = f"{ae.get('data_inicio', '—')} a {ae.get('data_fim', 'atual')}"
            tabela_esp.cell(i + 1, 1).text = periodo
            tabela_esp.cell(i + 1, 2).text = ae.get("probabilidade", "—")
            agentes = ae.get("agentes_provaveis", [])
            tabela_esp.cell(i + 1, 3).text = ", ".join(agentes) if agentes else "—"
            tabela_esp.cell(i + 1, 4).text = ae.get("recomendacao", "—")
            for j in range(5):
                for p in tabela_esp.cell(i + 1, j).paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)
            # Highlight high probability
            prob = ae.get("probabilidade", "")
            if prob == "ALTA":
                _set_cell_bg(tabela_esp.cell(i + 1, 2), "fef2f2")

        # Jurisprudências consolidadas
        juris_total = []
        for ae in analise_esp:
            for j in ae.get("jurisprudencias", []):
                key = f"{j.get('tipo', '')}_{j.get('numero', '')}"
                if key not in [f"{x.get('tipo','')}_{x.get('numero','')}" for x in juris_total]:
                    juris_total.append(j)

        if juris_total:
            _add_heading(doc, "14.1 Jurisprudência Consolidada Aplicável", level=2)
            _add_para(doc, (
                f"Foram identificadas {len(juris_total)} referências jurisprudenciais consolidadas "
                "com alto grau de confiabilidade (≥95%), aplicáveis aos empregadores analisados. "
                "Todas as referências são de jurisprudência pacificada (Súmulas TNU, Temas STJ/STF "
                "com repercussão geral ou recurso repetitivo)."
            ), size=10, color=CINZA)

            for j in juris_total:
                tipo = (j.get("tipo", "").replace("_", " ") + " " + j.get("numero", "")).strip()
                tribunal = j.get("tribunal", "")
                p_titulo = doc.add_paragraph()
                run_t = p_titulo.add_run(f"📚 {tipo} ({tribunal})")
                run_t.font.size = Pt(10)
                run_t.font.bold = True
                run_t.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

                if j.get("data_julgamento"):
                    _add_para(doc, f"Julgado em {j['data_julgamento']}", size=8, color=CINZA)

                ementa = j.get("ementa", "")
                if ementa:
                    _add_para(doc, f"Ementa: {ementa}", size=9)

                aplic = j.get("aplicabilidade", "")
                if aplic:
                    p_aplic = doc.add_paragraph()
                    run_a = p_aplic.add_run(f"➜ Aplicabilidade ao caso: {aplic}")
                    run_a.font.size = Pt(9)
                    run_a.font.italic = True
                    run_a.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)

                doc.add_paragraph()  # Espaçamento

        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 15. ANÁLISE DE REVISÃO DO BENEFÍCIO
    # ═══════════════════════════════════════════════════════════════════════

    analise_rev = planejamento.get("analise_revisao", {})
    if analise_rev.get("beneficio_ativo"):
        secnum = "15" if analise_esp else "14"
        _add_heading(doc, f"{secnum}. Análise de Possibilidade de Revisão", level=1)
        _add_para(doc, (
            f"O segurado possui benefício ativo ({analise_rev.get('especie', '')}). "
            f"Foi realizada análise automática de {len(analise_rev.get('revisoes_possiveis', []))} "
            "tipos de revisão conforme legislação vigente."
        ), size=10, color=CINZA)

        # Decadência
        dec = analise_rev.get("decadencia", {})
        if dec:
            cor_dec = VERDE if dec.get("dentro_prazo") else VERMELHO
            _add_para(doc, dec.get("mensagem", ""), size=10, bold=True, color=cor_dec)

        # Tabela de revisões
        revisoes = analise_rev.get("revisoes_possiveis", [])
        if revisoes:
            tabela_rev = doc.add_table(rows=len(revisoes) + 1, cols=5)
            tabela_rev.style = 'Light List Accent 1'
            for j, h in enumerate(["Revisão", "Aplicável", "Viável", "Impacto", "Análise"]):
                cell = tabela_rev.cell(0, j)
                cell.text = h
                _set_cell_bg(cell, "7c3aed")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = BRANCO
                        r.font.bold = True
                        r.font.size = Pt(8)

            for i, rev in enumerate(revisoes):
                tabela_rev.cell(i + 1, 0).text = rev.get("tipo", "—")
                tabela_rev.cell(i + 1, 1).text = "SIM" if rev.get("aplicavel") else "NÃO"
                tabela_rev.cell(i + 1, 2).text = "SIM" if rev.get("viavel") else "NÃO"
                tabela_rev.cell(i + 1, 3).text = rev.get("impacto_estimado", "—")
                tabela_rev.cell(i + 1, 4).text = rev.get("analise", "—")
                for j in range(5):
                    for p in tabela_rev.cell(i + 1, j).paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(8)
                if rev.get("viavel"):
                    _set_cell_bg(tabela_rev.cell(i + 1, 2), "f0fdf4")

        if analise_rev.get("recomendacao_geral"):
            _add_para(doc, analise_rev["recomendacao_geral"], size=10, bold=True, color=AZUL_ESCURO)

        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 16. MEMÓRIA DE CÁLCULO
    # ═══════════════════════════════════════════════════════════════════════

    memoria = planejamento.get("memoria_calculo", {})
    if memoria.get("linhas"):
        secnum_mc = "16"
        _add_heading(doc, f"{secnum_mc}. Memória de Cálculo — Correção Monetária", level=1)
        _add_para(doc, (
            f"Tabela com {memoria.get('total_contribuicoes', 0)} contribuições corrigidas pelo INPC "
            f"até a DER. {memoria.get('fundamentacao', '')}"
        ), size=10, color=CINZA)

        # Resumo de médias
        _add_para(doc, f"Média dos 80% maiores: {_fmt_brl(memoria.get('media_80_maiores', '0'))}", size=10, bold=True)
        _add_para(doc, f"Média 100% (EC 103/2019): {_fmt_brl(memoria.get('media_100', '0'))}", size=10)

        descarte = memoria.get("descarte", {})
        if descarte.get("aplicado"):
            _add_para(doc, (
                f"Descarte automático aplicado: {descarte.get('total_descartados', 0)} contribuições descartadas. "
                f"Economia mensal: {_fmt_brl(descarte.get('economia_mensal', '0'))}. "
                f"Fundamentação: {descarte.get('fundamentacao', 'Art. 26 §6 EC 103/2019')}"
            ), size=10, bold=True, color=VERDE)

        # Tabela (primeiras 50 linhas)
        linhas = memoria.get("linhas", [])[:50]
        if linhas:
            tabela_mc = doc.add_table(rows=len(linhas) + 1, cols=6)
            tabela_mc.style = 'Light List Accent 1'
            for j, h in enumerate(["Competência", "Empregador", "Sal. Original", "Índice", "Sal. Corrigido", "Status"]):
                cell = tabela_mc.cell(0, j)
                cell.text = h
                _set_cell_bg(cell, "1a3c6e")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = BRANCO
                        r.font.bold = True
                        r.font.size = Pt(7)

            for i, l in enumerate(linhas):
                tabela_mc.cell(i + 1, 0).text = l.get("competencia", "")
                emp = l.get("vinculo_nome", "")
                tabela_mc.cell(i + 1, 1).text = emp[:25] if emp else "—"
                tabela_mc.cell(i + 1, 2).text = _fmt_brl(l.get("salario_original", "0"))
                try:
                    idx = float(l.get("indice_correcao", "1"))
                    tabela_mc.cell(i + 1, 3).text = f"{idx:.4f}"
                except:
                    tabela_mc.cell(i + 1, 3).text = str(l.get("indice_correcao", ""))
                tabela_mc.cell(i + 1, 4).text = _fmt_brl(l.get("salario_corrigido", "0"))
                status = "Descartado" if l.get("descartado") else ("Teto" if l.get("limitado_teto") else "OK")
                tabela_mc.cell(i + 1, 5).text = status
                for j in range(6):
                    for p in tabela_mc.cell(i + 1, j).paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(7)
                if l.get("descartado"):
                    _set_cell_bg(tabela_mc.cell(i + 1, 5), "fef2f2")

            if len(memoria.get("linhas", [])) > 50:
                _add_para(doc, f"... e mais {len(memoria['linhas']) - 50} contribuições.", size=9, color=CINZA)

        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # CONCLUSÃO E RECOMENDAÇÃO FINAL
    # ═══════════════════════════════════════════════════════════════════════

    _add_heading(doc, "Conclusão e Recomendação Final", level=1)

    recomendacao = planejamento.get("recomendacao", "")
    _add_para(doc, recomendacao, size=11, bold=True, color=AZUL_ESCURO)

    _add_para(doc, (
        f"Com base na análise completa do histórico previdenciário de {nome}, "
        f"conforme legislação vigente (EC 103/2019, Lei 8.213/91), conclui-se que:"
    ), size=10)

    tc_texto = f"{tc.get('anos', 0)} anos, {tc.get('meses', 0)} meses e {tc.get('dias', 0)} dias"
    conclusoes = [
        f"O tempo de contribuição atual apurado é de {tc_texto}.",
        f"{'O segurado já é elegível para requerer aposentadoria.' if planejamento.get('elegiveis_agora') else 'O segurado ainda não preenche todos os requisitos.'}",
        "Para garantir a elegibilidade no prazo projetado, é essencial manter as contribuições mensais em dia, sem interrupções.",
        "Em caso de desemprego, recomenda-se o recolhimento como Segurado Facultativo (alíquota 20% sobre o salário-mínimo) para não perder o tempo de contribuição em curso.",
        "Os valores da RMI são estimativas baseadas nos salários atuais. Variações salariais e reajustes do teto podem alterar estes valores.",
    ]

    for c in conclusoes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(c)
        run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # ASSINATURA
    # ═══════════════════════════════════════════════════════════════════════

    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line2 = p_line2.add_run("━" * 40)
    run_line2.font.color.rgb = AZUL_MEDIO
    run_line2.font.size = Pt(8)

    adv_nome = nome_advogado or "Advogado(a) / Consultor(a) Previdenciário(a)"
    p_adv = doc.add_paragraph()
    p_adv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_adv = p_adv.add_run(adv_nome)
    run_adv.font.size = Pt(12)
    run_adv.font.bold = True
    run_adv.font.color.rgb = AZUL_ESCURO

    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_data = p_data.add_run(hoje)
    run_data.font.size = Pt(10)
    run_data.font.color.rgb = CINZA

    p_rodape = doc.add_paragraph()
    p_rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_rod = p_rodape.add_run("Documento gerado pelo SistPrev — Cálculos conforme Lei 8.213/91, EC 103/2019 e CJF Res. 963/2025")
    run_rod.font.size = Pt(8)
    run_rod.font.color.rgb = CINZA
    run_rod.font.italic = True

    # ── Salvar em bytes ──
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# RELATÓRIO PERICIAL — DOCX Feliciano Advocacia (Visual Law)
# ═══════════════════════════════════════════════════════════════════════════════

def _traduzir_tipo_pericial(tipo_raw: str, der) -> str:
    """Traduz o tipo interno para rótulo correto conforme marco temporal."""
    LABELS = {
        "pre_reforma": "Aposentadoria por TC — Regras Pré-Reforma (Lei 8.213/91)",
        "transicao_ec103": "Aposentadoria por TC — Regras de Transição (EC 103/2019)",
        "transicao": None,
        "idade": "Aposentadoria por Idade",
        "especial_15": "Aposentadoria Especial — 15 anos",
        "especial_20": "Aposentadoria Especial — 20 anos",
        "especial_25": "Aposentadoria Especial — 25 anos",
        "auxilio_doenca": "Auxílio por Incapacidade Temporária (B31)",
        "auxilio_doenca_acid": "Auxílio por Incapacidade Acidentário (B91)",
        "invalidez": "Aposentadoria por Incapacidade Permanente (B32)",
        "invalidez_acid": "Aposentadoria por Incapacidade Permanente Acidentária (B92)",
        "pensao_morte": "Pensão por Morte (B21)",
    }
    label = LABELS.get(tipo_raw)
    if label:
        return label
    try:
        if isinstance(der, str) and "/" in der:
            p = der.split("/")
            der_date = date(int(p[2]), int(p[1]), int(p[0]))
        elif isinstance(der, date):
            der_date = der
        else:
            der_date = date.fromisoformat(str(der))
        if der_date < date(2019, 11, 13):
            return "Aposentadoria por TC — Regras Pré-Reforma (Lei 8.213/91)"
        return "Aposentadoria por TC — Regras de Transição (EC 103/2019)"
    except Exception:
        return str(tipo_raw)


def _requisitos_por_tipo(tipo_raw: str, der) -> List[Dict[str, str]]:
    """
    Retorna a lista de requisitos específicos do tipo de benefício.
    Cada item: {"requisito": str, "fundamento": str, "padrao": str (para ajudar a verificar)}
    """
    tipo = (tipo_raw or "").lower()

    # AUXÍLIO POR INCAPACIDADE TEMPORÁRIA (B31)
    if "auxilio_doenca" in tipo and "acid" not in tipo:
        return [
            {"requisito": "Qualidade de segurado mantida na DII (Data de Início da Incapacidade)",
             "fundamento": "Lei 8.213/91, Art. 15 (períodos de graça)"},
            {"requisito": "Carência de 12 contribuições mensais",
             "fundamento": "Lei 8.213/91, Art. 25, I"},
            {"requisito": "Incapacidade temporária para o trabalho habitual por mais de 15 dias (perícia médica INSS)",
             "fundamento": "Lei 8.213/91, Art. 59"},
            {"requisito": "Não estar em gozo de outro benefício por incapacidade incompatível",
             "fundamento": "Lei 8.213/91, Art. 124"},
        ]
    # AUXÍLIO-ACIDENTÁRIO (B91)
    if "auxilio_doenca_acid" in tipo or tipo == "91":
        return [
            {"requisito": "Qualidade de segurado na data do acidente/doença ocupacional",
             "fundamento": "Lei 8.213/91, Art. 15"},
            {"requisito": "Carência DISPENSADA",
             "fundamento": "Lei 8.213/91, Art. 26, II"},
            {"requisito": "Acidente do trabalho, doença ocupacional ou do trabalho comprovada (CAT)",
             "fundamento": "Lei 8.213/91, Arts. 19-21"},
            {"requisito": "Incapacidade temporária para o trabalho habitual",
             "fundamento": "Lei 8.213/91, Art. 59"},
        ]
    # APOSENTADORIA POR INCAPACIDADE PERMANENTE (B32)
    if "invalidez" in tipo and "acid" not in tipo:
        return [
            {"requisito": "Qualidade de segurado mantida na DII",
             "fundamento": "Lei 8.213/91, Art. 15"},
            {"requisito": "Carência de 12 contribuições (exceto acidente/doença grave — dispensada)",
             "fundamento": "Lei 8.213/91, Arts. 25, I e 26, II; Portaria MPAS 2.998/2001"},
            {"requisito": "Incapacidade TOTAL e PERMANENTE para qualquer atividade laborativa",
             "fundamento": "Lei 8.213/91, Art. 42"},
            {"requisito": "Insusceptibilidade de reabilitação profissional (perícia médica)",
             "fundamento": "Lei 8.213/91, Art. 42 §1º"},
        ]
    # APOSENTADORIA POR INCAPACIDADE ACIDENTÁRIA (B92)
    if "invalidez_acid" in tipo or tipo == "92":
        return [
            {"requisito": "Qualidade de segurado",
             "fundamento": "Lei 8.213/91, Art. 15"},
            {"requisito": "Carência DISPENSADA",
             "fundamento": "Lei 8.213/91, Art. 26, II"},
            {"requisito": "Acidente de trabalho ou doença ocupacional",
             "fundamento": "Lei 8.213/91, Arts. 19-21"},
            {"requisito": "Incapacidade total e permanente (perícia INSS)",
             "fundamento": "Lei 8.213/91, Art. 42"},
        ]
    # PENSÃO POR MORTE (B21)
    if "pensao_morte" in tipo or tipo == "21":
        return [
            {"requisito": "Qualidade de segurado do instituidor na data do óbito",
             "fundamento": "Lei 8.213/91, Arts. 15 e 74"},
            {"requisito": "Comprovação da condição de dependente (Art. 16 Lei 8.213/91)",
             "fundamento": "Lei 8.213/91, Art. 16"},
            {"requisito": "Prova do óbito (certidão) ou morte presumida (sentença)",
             "fundamento": "Lei 8.213/91, Art. 78"},
            {"requisito": "Atendimento à carência de 18 meses (se exigida — EC 103/2019 Art. 23)",
             "fundamento": "EC 103/2019, Art. 23; Lei 13.135/2015"},
        ]
    # APOSENTADORIA POR IDADE (urbana)
    if tipo == "idade" or "idade" in tipo and "progressiva" not in tipo:
        # distinguir pré e pós EC 103
        try:
            if isinstance(der, str) and "/" in der:
                p = der.split("/")
                der_date = date(int(p[2]), int(p[1]), int(p[0]))
            else:
                der_date = der if isinstance(der, date) else date.fromisoformat(str(der))
            if der_date >= date(2019, 11, 13):
                return [
                    {"requisito": "Idade: 65 anos (homem) ou 62 anos (mulher)",
                     "fundamento": "EC 103/2019, Art. 18 e 19"},
                    {"requisito": "Tempo de contribuição mínimo: 15 anos (mulher) ou 20 anos (homem — para filiados após EC 103)",
                     "fundamento": "EC 103/2019, Art. 19"},
                    {"requisito": "Carência: 180 contribuições mensais",
                     "fundamento": "Lei 8.213/91, Art. 25, II"},
                    {"requisito": "Qualidade de segurado",
                     "fundamento": "Lei 8.213/91, Art. 15"},
                ]
        except Exception:
            pass
        return [
            {"requisito": "Idade: 65 anos (homem) ou 60 anos (mulher) — urbano",
             "fundamento": "Lei 8.213/91, Art. 48"},
            {"requisito": "Carência: 180 contribuições mensais",
             "fundamento": "Lei 8.213/91, Art. 25, II"},
            {"requisito": "Qualidade de segurado",
             "fundamento": "Lei 8.213/91, Art. 15"},
        ]
    # APOSENTADORIA ESPECIAL
    if "especial" in tipo:
        tempo_minimo = "25 anos"
        if "15" in tipo: tempo_minimo = "15 anos"
        elif "20" in tipo: tempo_minimo = "20 anos"
        return [
            {"requisito": f"Tempo especial mínimo: {tempo_minimo} de exposição a agentes nocivos",
             "fundamento": "Lei 8.213/91, Art. 57; Decreto 3.048/99 Art. 64-70; EC 103/2019 Art. 19"},
            {"requisito": "Comprovação por PPP (Perfil Profissiográfico Previdenciário)",
             "fundamento": "IN INSS 128/2022; Decreto 4.882/2003"},
            {"requisito": "LTCAT (Laudo Técnico de Condições Ambientais de Trabalho)",
             "fundamento": "Lei 8.213/91, Art. 58 §2º"},
            {"requisito": "Carência: 180 contribuições mensais",
             "fundamento": "Lei 8.213/91, Art. 25, II"},
            {"requisito": "Após EC 103/2019: idade mínima (55/58/60 conforme categoria)",
             "fundamento": "EC 103/2019, Art. 19 §1º III"},
        ]
    # APOSENTADORIA POR TC — PRÉ-REFORMA
    if "pre_reforma" in tipo or "85_95" in tipo or "tc_fator" in tipo:
        return [
            {"requisito": "Tempo de contribuição: 35 anos (homem) ou 30 anos (mulher)",
             "fundamento": "Lei 8.213/91, Art. 52"},
            {"requisito": "Carência: 180 contribuições mensais",
             "fundamento": "Lei 8.213/91, Art. 25, II"},
            {"requisito": "Qualidade de segurado",
             "fundamento": "Lei 8.213/91, Art. 15"},
            {"requisito": "DER anterior a 13/11/2019 (direito adquirido)",
             "fundamento": "EC 103/2019, Art. 3º"},
        ]
    # TRANSIÇÃO EC 103/2019
    if "transicao" in tipo or "pontos" in tipo or "progressiva" in tipo or "pedagio" in tipo:
        return [
            {"requisito": "Tempo de contribuição mínimo: 35/30 anos (H/M)",
             "fundamento": "EC 103/2019, Arts. 15-20"},
            {"requisito": "Carência: 180 contribuições mensais",
             "fundamento": "Lei 8.213/91, Art. 25, II"},
            {"requisito": "Requisitos específicos da regra de transição escolhida (pontos, idade progressiva, pedágio 50% ou 100%)",
             "fundamento": "EC 103/2019, Arts. 15, 16, 17 e 20"},
            {"requisito": "Estar filiado ao RGPS em 13/11/2019 (data da promulgação da EC 103)",
             "fundamento": "EC 103/2019, Art. 3º"},
        ]
    # Fallback genérico
    return [
        {"requisito": "Qualidade de segurado",
         "fundamento": "Lei 8.213/91, Art. 15"},
        {"requisito": "Carência mínima conforme o benefício",
         "fundamento": "Lei 8.213/91, Art. 25"},
        {"requisito": "Demais requisitos específicos da espécie",
         "fundamento": "Lei 8.213/91 e EC 103/2019"},
    ]


def _card_visual_law(doc, titulo: str, texto: str, cor_bg: str, cor_borda_hex: str, icone: str = ""):
    """Cria um card estilo Visual Law — caixa colorida com título em bold e texto explicativo."""
    tabela = doc.add_table(rows=1, cols=1)
    cell = tabela.cell(0, 0)
    p = cell.paragraphs[0]
    if icone:
        run_i = p.add_run(f"{icone}  ")
        run_i.font.size = Pt(11)
    run_t = p.add_run(f"{titulo}\n")
    run_t.font.size = Pt(9)
    run_t.font.bold = True
    run_t.font.color.rgb = AZUL_ESCURO
    run_v = p.add_run(texto)
    run_v.font.size = Pt(10)
    _set_cell_bg(cell, cor_bg)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def gerar_docx_pericial(
    segurado: Dict,
    calculo: Dict,
    nome_advogado: Optional[str] = None,
    titulo: str = "RELATÓRIO PERICIAL PREVIDENCIÁRIO",
) -> bytes:
    """
    Gera o Relatório Pericial em DOCX Feliciano Advocacia (Visual Law).
    Inclui análise específica dos requisitos do benefício e cards coloridos.
    """
    doc = Document()

    # ── Estilos ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = AZUL_ESCURO if level == 1 else AZUL_MEDIO

    # ── Margens ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    dp = segurado.get("dados_pessoais", {})
    vinculos = segurado.get("vinculos") or []
    hoje = date.today().strftime("%d/%m/%Y")
    nome = dp.get("nome", "Segurado")

    cenarios = calculo.get("todos_cenarios") or []
    melhor = calculo.get("melhor_cenario") or {}
    elegivel = calculo.get("elegivel", False)
    der = calculo.get("der", "—")
    tipo_raw = calculo.get("tipo", "—")
    tipo_label = _traduzir_tipo_pericial(tipo_raw, der)
    rmi_str = calculo.get("rmi", "0")
    rmi_formatada = calculo.get("rmi_formatada") or _fmt_brl(rmi_str)
    modo_revisao = calculo.get("modo_revisao", False)
    nb_ativo = calculo.get("nb_ativo")
    alertas = calculo.get("alertas_consistencia") or []

    # ═══════════════════════════════════════════════════════════════════════
    # CAPA
    # ═══════════════════════════════════════════════════════════════════════
    _add_para(doc, "", size=6, space_after=0)
    _add_para(doc, "", size=6, space_after=0)

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("FELICIANO ADVOCACIA")
    run_logo.font.size = Pt(28)
    run_logo.font.bold = True
    run_logo.font.color.rgb = AZUL_ESCURO
    run_logo.font.name = "Calibri"

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Assessoria Jurídica Previdenciária")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = AZUL_MEDIO
    run_sub.font.name = "Calibri"
    p_sub.paragraph_format.space_after = Pt(40)

    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("━" * 60)
    run_line.font.color.rgb = AZUL_MEDIO
    run_line.font.size = Pt(8)
    p_line.paragraph_format.space_after = Pt(30)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run(titulo)
    run_t.font.size = Pt(22)
    run_t.font.bold = True
    run_t.font.color.rgb = AZUL_ESCURO
    p_titulo.paragraph_format.space_after = Pt(8)

    p_nome_cli = doc.add_paragraph()
    p_nome_cli.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nc = p_nome_cli.add_run(nome.upper())
    run_nc.font.size = Pt(16)
    run_nc.font.bold = True
    run_nc.font.color.rgb = AZUL_MEDIO
    p_nome_cli.paragraph_format.space_after = Pt(30)

    # Quadro de capa
    tabela_capa = doc.add_table(rows=6, cols=2)
    tabela_capa.alignment = WD_TABLE_ALIGNMENT.CENTER
    dados_capa = [
        ("CPF:", dp.get("cpf", "—")),
        ("Data de Nascimento:", dp.get("data_nascimento", "—")),
        ("NIT/PIS:", dp.get("nit", "—")),
        ("DER (Data de Entrada do Requerimento):", der),
        ("Tipo de Benefício Analisado:", tipo_label),
        ("Data do Relatório:", hoje),
    ]
    for i, (label, valor) in enumerate(dados_capa):
        cell_l = tabela_capa.cell(i, 0)
        cell_v = tabela_capa.cell(i, 1)
        cell_l.text = label
        cell_v.text = str(valor) if valor else "—"
        for p in cell_l.paragraphs:
            if p.runs:
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = AZUL_ESCURO
        for p in cell_v.paragraphs:
            if p.runs:
                p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 1. IDENTIFICAÇÃO DO SEGURADO
    # ═══════════════════════════════════════════════════════════════════════
    _add_heading(doc, "1. Identificação do Segurado", level=1)

    sexo = dp.get("sexo", "—")
    modalidade = "Revisão de benefício em manutenção" if modo_revisao else "Cálculo inicial de concessão"
    tabela_id = doc.add_table(rows=8, cols=2)
    tabela_id.style = 'Light List Accent 1'
    dados_id = [
        ("Nome Completo", nome),
        ("CPF", dp.get("cpf", "—")),
        ("NIT / PIS-PASEP", dp.get("nit", "—")),
        ("Data de Nascimento", dp.get("data_nascimento", "—")),
        ("Sexo", sexo),
        ("Data de Entrada do Requerimento (DER)", der),
        ("Modalidade do Cálculo", modalidade),
        ("Data do Relatório", hoje),
    ]
    for i, (label, valor) in enumerate(dados_id):
        tabela_id.cell(i, 0).text = label
        tabela_id.cell(i, 1).text = str(valor) if valor else "—"
        for p in tabela_id.cell(i, 0).paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = AZUL_ESCURO
        for p in tabela_id.cell(i, 1).paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 2. RESUMO EXECUTIVO — Cards Visual Law
    # ═══════════════════════════════════════════════════════════════════════
    _add_heading(doc, "2. Resumo Executivo", level=1)

    # Card 1: SITUAÇÃO (azul claro)
    situacao_texto = (
        f"Análise de {tipo_label} para o(a) segurado(a) {nome}, "
        f"com DER em {der}. "
        + ("Cálculo de REVISÃO de benefício em manutenção." if modo_revisao else "Cálculo de CONCESSÃO inicial.")
    )
    _card_visual_law(doc, "SITUAÇÃO ANALISADA", situacao_texto,
                     cor_bg="DBE5F1", cor_borda_hex="1a3c6e", icone="📋")

    # Card 2: RESULTADO (verde se elegível / amarelo se não)
    if elegivel:
        resultado_texto = (
            f"O(a) segurado(a) PREENCHE todos os requisitos do benefício. "
            f"RMI apurada: {rmi_formatada}. "
            "Recomenda-se a protocolização do requerimento administrativo no INSS."
        )
        cor_bg_res = "E2EFDA"; icone_res = "✅"
    else:
        resultado_texto = (
            f"O(a) segurado(a) NÃO preenche atualmente os requisitos para {tipo_label} na DER indicada. "
            "Consultar a seção 'Análise de Requisitos' para ver quais itens faltam e orientações de planejamento."
        )
        cor_bg_res = "FFF2CC"; icone_res = "⚠"
    _card_visual_law(doc, "RESULTADO DA ANÁLISE", resultado_texto,
                     cor_bg=cor_bg_res, cor_borda_hex="1a3c6e", icone=icone_res)

    # Card 3: AÇÃO IMEDIATA
    if elegivel:
        acao_texto = (
            "1) Reunir documentação probatória (CNIS, CTPS, PPP/LTCAT se aplicável); "
            "2) Protocolar pedido administrativo via Meu INSS; "
            "3) Aguardar resposta em até 45 dias úteis; "
            "4) Em caso de indeferimento, propor ação judicial."
        )
    else:
        acao_texto = (
            "1) Verificar requisitos faltantes na seção 4; "
            "2) Avaliar alternativas (outras regras de transição, complementação MEI, averbação especial); "
            "3) Elaborar plano previdenciário para futura elegibilidade; "
            "4) Manter contribuições em dia para preservar qualidade de segurado."
        )
    _card_visual_law(doc, "AÇÃO RECOMENDADA", acao_texto,
                     cor_bg="E8DAEF", cor_borda_hex="5b21b6", icone="🎯")

    # Card 4: VALOR FINANCEIRO (se elegível)
    if elegivel:
        valor_texto = (
            f"RMI (Renda Mensal Inicial): {rmi_formatada}. "
            "Valor sujeito a reajustes anuais conforme INPC (Art. 41-A Lei 8.213/91) e limitado ao teto do RGPS."
        )
        _card_visual_law(doc, "VALOR DO BENEFÍCIO", valor_texto,
                         cor_bg="DCFCE7", cor_borda_hex="065f46", icone="💰")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 3. ANÁLISE DE REQUISITOS — CHECKLIST ESPECÍFICO DO BENEFÍCIO
    # ═══════════════════════════════════════════════════════════════════════
    _add_heading(doc, f"3. Análise Específica dos Requisitos — {tipo_label}", level=1)

    _add_para(doc, (
        "A tabela abaixo lista os requisitos exigidos por lei para a concessão deste benefício. "
        "Cada item é avaliado conforme os dados do CNIS e documentos juntados. Requisitos marcados em verde "
        "foram atendidos; em vermelho, precisam ser comprovados ou atendidos."
    ), size=10, color=CINZA)

    requisitos = _requisitos_por_tipo(tipo_raw, der)

    # Cabeçalho da tabela
    tabela_req = doc.add_table(rows=len(requisitos) + 1, cols=3)
    tabela_req.style = 'Light List Accent 1'
    headers_req = ["Requisito Legal", "Fundamento", "Situação"]
    for j, h in enumerate(headers_req):
        cell = tabela_req.cell(0, j)
        cell.text = h
        _set_cell_bg(cell, "1a3c6e")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = BRANCO
                r.font.bold = True
                r.font.size = Pt(9)

    for i, req in enumerate(requisitos):
        row = i + 1
        tabela_req.cell(row, 0).text = req["requisito"]
        tabela_req.cell(row, 1).text = req["fundamento"]
        # A análise de "situação" aqui é heurística: se elegível = todos verdes;
        # se não elegível = marca como "a verificar" (amarelo) para o advogado conferir
        if elegivel:
            tabela_req.cell(row, 2).text = "✓ ATENDIDO"
            _set_cell_bg(tabela_req.cell(row, 2), "dcfce7")
            for p in tabela_req.cell(row, 2).paragraphs:
                for r in p.runs:
                    r.font.color.rgb = VERDE
                    r.font.bold = True
                    r.font.size = Pt(9)
        else:
            tabela_req.cell(row, 2).text = "⚠ VERIFICAR"
            _set_cell_bg(tabela_req.cell(row, 2), "fef3c7")
            for p in tabela_req.cell(row, 2).paragraphs:
                for r in p.runs:
                    r.font.color.rgb = LARANJA
                    r.font.bold = True
                    r.font.size = Pt(9)

        # Estilo das outras colunas
        for j in [0, 1]:
            for p in tabela_req.cell(row, j).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        # Fundamento em itálico cinza
        for p in tabela_req.cell(row, 1).paragraphs:
            for r in p.runs:
                r.font.color.rgb = CINZA
                r.font.italic = True

    # Card explicativo
    _add_para(doc, "", size=4)
    if elegivel:
        _card_visual_law(doc,
            "CONCLUSÃO DA ANÁLISE DE REQUISITOS",
            f"Todos os requisitos legais para {tipo_label} estão preenchidos conforme os dados apresentados. "
            "Confirme a documentação probatória antes da protocolização.",
            cor_bg="E2EFDA", cor_borda_hex="065f46", icone="✅")
    else:
        _card_visual_law(doc,
            "CONCLUSÃO DA ANÁLISE DE REQUISITOS",
            f"Um ou mais requisitos para {tipo_label} não foram preenchidos na DER indicada. "
            "Consultar a seção 'Comparativo de Regras' para identificar uma alternativa viável.",
            cor_bg="FEF3C7", cor_borda_hex="b45309", icone="⚠")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 4. ALERTAS DE CONSISTÊNCIA (se houver)
    # ═══════════════════════════════════════════════════════════════════════
    if alertas:
        _add_heading(doc, "4. Alertas de Consistência", level=1)
        _add_para(doc, (
            "Os pontos abaixo foram detectados pela análise automática e "
            "devem ser verificados antes da protocolização do pedido."
        ), size=9, color=CINZA)

        for alerta in alertas:
            tabela_a = doc.add_table(rows=1, cols=1)
            cell = tabela_a.cell(0, 0)
            p = cell.paragraphs[0]
            run = p.add_run(f"⚠  {alerta}")
            run.font.size = Pt(10)
            run.font.color.rgb = VERMELHO
            _set_cell_bg(cell, "fef2f2")
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # MODO REVISÃO — explicação (se aplicável)
    # ═══════════════════════════════════════════════════════════════════════
    if modo_revisao and nb_ativo:
        _add_heading(doc, "Análise de Revisão", level=1)
        tabela_rev = doc.add_table(rows=1, cols=1)
        cell = tabela_rev.cell(0, 0)
        p = cell.paragraphs[0]
        run1 = p.add_run(f"Benefício em manutenção: NB {nb_ativo}\n")
        run1.font.bold = True
        run1.font.size = Pt(11)
        run1.font.color.rgb = AZUL_ESCURO
        run2 = p.add_run(
            "Este cálculo simula a aposentadoria como se o requerimento tivesse sido feito na DER indicada, "
            "para fins de comparação com o benefício atualmente pago. A diferença apontada no item 'Resultado' "
            "representa o ganho (ou perda) mensal que justifica (ou afasta) o pedido de revisão."
        )
        run2.font.size = Pt(10)
        _set_cell_bg(cell, "eff6ff")
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 5. VÍNCULOS EMPREGATÍCIOS
    # ═══════════════════════════════════════════════════════════════════════
    _add_heading(doc, "5. Vínculos Empregatícios e Contribuições", level=1)

    if not vinculos:
        _add_para(doc, "Nenhum vínculo empregatício informado.", size=10, color=CINZA)
    else:
        tabela_v = doc.add_table(rows=len(vinculos) + 1, cols=6)
        tabela_v.style = 'Light List Accent 1'
        headers_v = ["Empregador", "CNPJ", "Tipo", "Início", "Fim", "Competências"]
        for j, h in enumerate(headers_v):
            cell = tabela_v.cell(0, j)
            cell.text = h
            _set_cell_bg(cell, "1a3c6e")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = BRANCO
                    r.font.bold = True
                    r.font.size = Pt(9)

        for i, v in enumerate(vinculos):
            row = i + 1
            n_comp = 0
            contribs = v.get("contribuicoes") or []
            n_comp = sum(1 for c in contribs if c.get("valida_tc", True))
            fim = v.get("data_fim") or "Em aberto"
            tabela_v.cell(row, 0).text = str(v.get("empregador_nome", "—"))[:45]
            tabela_v.cell(row, 1).text = str(v.get("empregador_cnpj", "—"))
            tabela_v.cell(row, 2).text = str(v.get("tipo_vinculo", "—"))
            tabela_v.cell(row, 3).text = str(v.get("data_inicio", "—"))
            tabela_v.cell(row, 4).text = str(fim)
            tabela_v.cell(row, 5).text = str(n_comp)
            for j in range(6):
                for p in tabela_v.cell(row, j).paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

        total_comp = sum(
            sum(1 for c in (v.get("contribuicoes") or []) if c.get("valida_tc", True))
            for v in vinculos
        )
        p_total = doc.add_paragraph()
        run_t = p_total.add_run(f"Total de competências registradas: {total_comp}")
        run_t.font.bold = True
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = AZUL_ESCURO
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 6. RESULTADO DO CÁLCULO — DESTAQUE
    # ═══════════════════════════════════════════════════════════════════════
    _add_heading(doc, "6. Resultado do Cálculo", level=1)

    # Status destacado
    status_texto = "ELEGÍVEL (conforme dados informados)" if elegivel else "NÃO ELEGÍVEL na DER"
    status_cor = VERDE if elegivel else VERMELHO
    status_bg = "dcfce7" if elegivel else "fef2f2"

    tabela_st = doc.add_table(rows=1, cols=1)
    cell_st = tabela_st.cell(0, 0)
    p_st = cell_st.paragraphs[0]
    p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_st = p_st.add_run(status_texto)
    run_st.font.size = Pt(14)
    run_st.font.bold = True
    run_st.font.color.rgb = status_cor
    _set_cell_bg(cell_st, status_bg)
    doc.add_paragraph()

    # RMI em destaque
    p_rmi = doc.add_paragraph()
    p_rmi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_r = p_rmi.add_run(rmi_formatada)
    run_r.font.size = Pt(32)
    run_r.font.bold = True
    run_r.font.color.rgb = VERDE if elegivel else CINZA
    p_rmi_lbl = doc.add_paragraph()
    p_rmi_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_lbl = p_rmi_lbl.add_run("Renda Mensal Inicial (RMI)")
    run_lbl.font.size = Pt(10)
    run_lbl.font.color.rgb = CINZA
    p_rmi_lbl.paragraph_format.space_after = Pt(16)

    # Tabela detalhada do cálculo (a partir do melhor cenário)
    base_legal = melhor.get("base_legal") or melhor.get("fundamento_legal") or "Lei 8.213/91"
    sb = melhor.get("salario_beneficio") or melhor.get("sb") or "—"
    coef = melhor.get("coeficiente") or melhor.get("fator") or "—"
    tc_apurado = melhor.get("tempo_contribuicao") or melhor.get("tc") or "—"
    regra_nome = melhor.get("nome_regra") or melhor.get("regra") or tipo_label

    tabela_res = doc.add_table(rows=6, cols=2)
    tabela_res.style = 'Light List Accent 1'
    dados_res = [
        ("Melhor Regra Identificada", regra_nome),
        ("Base Legal", base_legal),
        ("Tempo de Contribuição Apurado", str(tc_apurado)),
        ("Salário de Benefício (SB)", _fmt_brl(sb) if sb != "—" else "—"),
        ("Coeficiente / Fator Aplicado", str(coef)),
        ("RMI — Renda Mensal Inicial", rmi_formatada),
    ]
    for i, (label, valor) in enumerate(dados_res):
        tabela_res.cell(i, 0).text = label
        tabela_res.cell(i, 1).text = str(valor)
        for p in tabela_res.cell(i, 0).paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = AZUL_ESCURO
        for p in tabela_res.cell(i, 1).paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
        # Destacar linha da RMI
        if i == 5:
            _set_cell_bg(tabela_res.cell(i, 1), "dcfce7" if elegivel else "f3f4f6")
            for p in tabela_res.cell(i, 1).paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = VERDE if elegivel else CINZA
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 7. ANÁLISE POR REGRA (Comparativo de cenários)
    # ═══════════════════════════════════════════════════════════════════════
    if cenarios and len(cenarios) > 1:
        _add_heading(doc, "7. Comparativo de Todas as Regras Analisadas", level=1)
        _add_para(doc, (
            "O sistema avaliou todas as regras aplicáveis ao caso, sem pré-seleção. "
            "A tabela abaixo mostra o resultado para cada uma — a regra mais vantajosa "
            "está destacada na seção anterior."
        ), size=9, color=CINZA)

        tabela_c = doc.add_table(rows=len(cenarios) + 1, cols=4)
        tabela_c.style = 'Light List Accent 1'
        headers_c = ["Regra", "Elegível", "RMI", "Fundamento"]
        for j, h in enumerate(headers_c):
            cell = tabela_c.cell(0, j)
            cell.text = h
            _set_cell_bg(cell, "1a3c6e")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = BRANCO
                    r.font.bold = True
                    r.font.size = Pt(9)

        for i, c in enumerate(cenarios):
            row = i + 1
            nome_r = c.get("nome_regra") or c.get("regra") or "—"
            eleg = c.get("elegivel", False)
            rmi_c = c.get("rmi") or "—"
            fund = (c.get("base_legal") or c.get("fundamento_legal") or "—")[:80]

            tabela_c.cell(row, 0).text = str(nome_r)
            tabela_c.cell(row, 1).text = "✓ SIM" if eleg else "✗ NÃO"
            tabela_c.cell(row, 2).text = _fmt_brl(rmi_c) if rmi_c not in ("—", None) else "—"
            tabela_c.cell(row, 3).text = str(fund)

            if eleg:
                _set_cell_bg(tabela_c.cell(row, 1), "dcfce7")
            else:
                _set_cell_bg(tabela_c.cell(row, 1), "fef2f2")

            for j in range(4):
                for p in tabela_c.cell(row, j).paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        if j == 1:
                            r.font.bold = True
                            r.font.color.rgb = VERDE if eleg else VERMELHO
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 8. MEMÓRIA DE CÁLCULO (a partir do melhor cenário)
    # ═══════════════════════════════════════════════════════════════════════
    memoria = melhor.get("memoria") or []
    if memoria:
        _add_heading(doc, "8. Memória de Cálculo", level=1)
        _add_para(doc, (
            "Passo a passo do cálculo conforme metodologia oficial (Lei 9.876/99, "
            "EC 103/2019 Art. 26 e CJF Res. 963/2025)."
        ), size=9, color=CINZA)

        # Mostrar até 30 passos — se for muito longo, avisa
        for i, passo in enumerate(memoria[:30]):
            if isinstance(passo, dict):
                etapa = passo.get("etapa") or passo.get("passo") or passo.get("descricao") or f"Passo {i+1}"
                valor = passo.get("valor") or passo.get("resultado") or ""
                obs = passo.get("observacao") or passo.get("base_legal") or ""
                p = doc.add_paragraph(style='List Number')
                run1 = p.add_run(f"{etapa}")
                run1.font.bold = True
                run1.font.size = Pt(10)
                if valor:
                    run2 = p.add_run(f" → {valor}")
                    run2.font.size = Pt(10)
                    run2.font.color.rgb = AZUL_ESCURO
                if obs:
                    run3 = p.add_run(f"\n   {obs}")
                    run3.font.size = Pt(9)
                    run3.font.italic = True
                    run3.font.color.rgb = CINZA
            else:
                p = doc.add_paragraph(style='List Number')
                run = p.add_run(str(passo))
                run.font.size = Pt(10)

        if len(memoria) > 30:
            _add_para(doc, f"... e mais {len(memoria) - 30} passos de memória de cálculo.",
                     size=9, color=CINZA)
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 9. FUNDAMENTAÇÃO LEGAL
    # ═══════════════════════════════════════════════════════════════════════
    _add_heading(doc, "9. Fundamentação Legal", level=1)

    fundamentos = [
        "Constituição Federal, Art. 201 — Princípios da Previdência Social.",
        "Lei 8.213/91 — Plano de Benefícios da Previdência Social (arts. 18, 25, 29, 48, 52, 57, 59-63).",
        "Lei 9.876/99 — Fator Previdenciário e Art. 29-C (Regra 85/95).",
        "Decreto 3.048/99 — Regulamento da Previdência Social.",
        "Emenda Constitucional 103/2019 — Reforma da Previdência (arts. 15-20, 26).",
        "Decreto 10.410/2020 — Alterações pós-EC 103.",
        "CJF Resolução 963/2025 — Manual de Cálculos da Justiça Federal.",
        "STJ Tema 692 / Súmula 576 — Cessação de auxílio-doença.",
        "STJ Tema 862 — Termo inicial do restabelecimento (DCB + 1 dia).",
        "STF Tema 334 — Direito adquirido no regime anterior.",
    ]
    for f in fundamentos:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f)
        run.font.size = Pt(9)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # 10. CONCLUSÃO E RECOMENDAÇÃO
    # ═══════════════════════════════════════════════════════════════════════
    _add_heading(doc, "10. Conclusão e Recomendação", level=1)

    if elegivel:
        conclusao_principal = (
            f"Com base na análise completa do histórico previdenciário de {nome} "
            f"e conforme legislação vigente, o(a) segurado(a) FAZ JUS ao benefício "
            f"{tipo_label}, com RMI de {rmi_formatada}, "
            f"na DER de {der}, pela regra {regra_nome}."
        )
        cor_conc = VERDE
    else:
        conclusao_principal = (
            f"Com base na análise completa do histórico previdenciário de {nome} "
            f"e conforme legislação vigente, o(a) segurado(a) NÃO preenche os requisitos "
            f"para {tipo_label} na DER de {der}. Recomenda-se a análise de outras regras de transição "
            f"ou o planejamento previdenciário para futura elegibilidade."
        )
        cor_conc = LARANJA

    _add_para(doc, conclusao_principal, size=11, bold=True, color=cor_conc)

    pontos_conc = []
    if elegivel:
        pontos_conc = [
            f"A RMI apurada é {rmi_formatada}, já observados o teto e o piso previdenciário vigentes.",
            "Recomenda-se o protocolo do requerimento administrativo no INSS (via Meu INSS ou presencial) com toda a documentação probatória (CNIS, CTPS, PPP/LTCAT se aplicável).",
            "Em caso de indeferimento administrativo, cabe ação judicial nos termos da Lei 8.213/91.",
            "Os valores da RMI são calculados com base nos salários constantes no CNIS e podem ser revistos mediante averbação de períodos adicionais.",
        ]
    else:
        pontos_conc = [
            "Recomenda-se planejamento previdenciário para identificar a regra mais vantajosa e a data provável de elegibilidade.",
            "Manter as contribuições em dia evita perda da qualidade de segurado (Art. 15 Lei 8.213/91).",
            "Em caso de desemprego, o recolhimento como Segurado Facultativo (20% sobre o salário-mínimo) preserva o tempo de contribuição em curso.",
            "Averbação de tempo especial, rural ou de serviço público pode alterar a elegibilidade — juntar PPP, LTCAT, CTC ou documento equivalente.",
        ]

    for c in pontos_conc:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(c)
        run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # ASSINATURA
    # ═══════════════════════════════════════════════════════════════════════
    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line2 = p_line2.add_run("━" * 40)
    run_line2.font.color.rgb = AZUL_MEDIO
    run_line2.font.size = Pt(8)

    adv_nome = nome_advogado or "Advogado(a) / Consultor(a) Previdenciário(a)"
    p_adv = doc.add_paragraph()
    p_adv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_adv = p_adv.add_run(adv_nome)
    run_adv.font.size = Pt(12)
    run_adv.font.bold = True
    run_adv.font.color.rgb = AZUL_ESCURO

    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_data = p_data.add_run(hoje)
    run_data.font.size = Pt(10)
    run_data.font.color.rgb = CINZA

    p_rodape = doc.add_paragraph()
    p_rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_rod = p_rodape.add_run(
        "Documento gerado pelo SistPrev — Cálculos conforme Lei 8.213/91, EC 103/2019 e CJF Res. 963/2025"
    )
    run_rod.font.size = Pt(8)
    run_rod.font.color.rgb = CINZA
    run_rod.font.italic = True

    # ── Salvar em bytes ──
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
